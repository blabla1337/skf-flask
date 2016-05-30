# -*- coding: utf-8 -*-
"""
    Security Knowledge Framework is an expert system application
    that uses OWASP Application Security Verification Standard, code examples,
    helps developers in pre-development and post-development.
    Copyright (C) 2015  Glenn ten Cate, Riccardo ten Cate

    This program is free software: you can redistribute it and/or modify
    it under the terms of the GNU Affero General Public License as
    published by the Free Software Foundation, either version 3 of the
    License, or (at your option) any later version.

    This program is distributed in the hope that it will be useful,
    but WITHOUT ANY WARRANTY; without even the implied warranty of
    MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
    GNU Affero General Public License for more details.

    You should have received a copy of the GNU Affero General Public License
    along with this program. If not, see <http://www.gnu.org/licenses/>.
"""

import contextlib, traceback
import os, markdown, datetime, string, base64, re, sys, re, requests, mimetypes, smtplib
from OpenSSL import SSL, rand
from docx import Document
from BeautifulSoup import BeautifulSoup
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from functools import wraps
from sqlite3 import dbapi2 as sqlite3
from flask_bcrypt import Bcrypt
from flask import Flask, request, session, g, redirect, url_for, \
     render_template, flash, Markup, make_response



# create the application
app = Flask(__name__)
app.jinja_env.add_extension("jinja2.ext.loopcontrols")

"""Set up bcrypt for passwords encrypting"""
bcrypt = Bcrypt(app)


def add_response_headers(headers={}):
    """This decorator adds the headers passed in to the response"""
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            resp = make_response(f(*args, **kwargs))
            h = resp.headers
            for header, value in headers.items():
                h[header] = value
            return resp
        return decorated_function
    return decorator


def security(f):
    """This decorator passes multiple security headers and checks log file to block users"""
    # return add_response_headers({'X-Frame-Options': 'deny', 'X-XSS-Protection': '1', 'X-Content-Type-Options': 'nosniff', 'Cache-Control': 'no-store, no-cache','Strict-Transport-Security': 'max-age=16070400; includeSubDomains', 'Server': 'Security Knowledge Framework'})(f)
    return add_response_headers({'X-Frame-Options': 'deny', 'X-XSS-Protection': '1', 'X-Content-Type-Options': 'nosniff', 'Cache-Control': 'no-store, no-cache', 'Server': 'Security Knowledge Framework'})(f)


def check_token():
    """Checks the submitted CSRF token"""
    if not session.get('csrf_token') == request.form['csrf_token']:
        raise Exception("User supplied CSRF token not valid")


def generate_pass():
    chars = string.letters + string.digits + '+/'
    assert 256 % len(chars) == 0  # non-biased later modulo
    PWD_LEN = 12
    generated_pass = ''.join(chars[ord(c) % len(chars)] for c in os.urandom(PWD_LEN))
    return generated_pass


def random_token(tokenBytes):
    #Create random token
    rand.cleanup()
    Random_token_raw = rand.bytes(int(tokenBytes))
    Random_token = base64.b64encode(Random_token_raw)
    result = re.sub("==", "", Random_token)
    return result


def log(message, value, threat):
    """Create log file and write events triggerd by the user
    The variables: message can be everything, value contains FAIL or SUCCESS and threat LOW MEDIUM HIGH"""
    now = datetime.datetime.now()
    dateLog = now.strftime("%Y-%m")
    dateTime = now.strftime("%Y-%m-%d %H:%M")
    ip = request.remote_addr
    fullmessage = dateTime +' '+ message +' ' + ' ' + value + ' ' + threat + ' ' + ip
    sys.stderr.write(fullmessage)
    try:
        file = open('logs/'+dateLog+'.txt', 'a+')
    except IOError:
        # If not exists, create the file
        file = open('logs/'+dateLog+'.txt', 'w+')
    file.write(fullmessage + "\r\n")
    file.close()


def blockUsers():
    """Check the log file and based on the FAIL items block a user"""
    dateLog  = datetime.datetime.now().strftime("%Y-%m")
    count = 0
    try:
        read = open(os.path.join(app.root_path, 'logs/'+dateLog+'.txt'), 'a+')
    except IOError:
        # If not exists, create the file
        read = open(os.path.join(app.root_path, 'logs/'+dateLog+'.txt'), 'w+')
    for line in read:
        match = re.search('FAIL', line)
        # If-statement after search() tests if it succeeded
        if match:
            count += 1
            str(count)
            if count > 11:
                sys.exit('Due to to many FAILED logs in your logging file we have the suspicion your application has been under attack by hackers. Please check your log files to validate and take repercussions. After validation clear your log or simply change the FAIL items to another value.')


def whiteList(range, value, countLevel):
    match = re.findall(range, value)
    if match:
        return True
    else:
        raise Exception("User supplied value not in the range " + range)


def valAlphaNum(value, countLevel):
    return whiteList(r'^([ a-zA-Z0-9_.-]*)$', value, countLevel)


def valNum(value, countLevel):
    return whiteList(r'^([0-9]+)$', value, countLevel)


def valBool(value, countLevel):
    return whiteList(r'^(true|false)$', value, countLevel)


#secret key for flask internal session use
rand.cleanup()
secret_key = rand.bytes(512)

mimetypes.add_type('image/svg+xml', '.svg')
bindaddr = '127.0.0.1';

# Load default config and override config from an environment variable
# You can also replace password with static password:  PASSWORD='pass!@#example'
app.config.update(dict(
    DATABASE=os.path.join(app.root_path, 'skf.db'),
    DEBUG=False,
    SECRET_KEY=secret_key,
    SESSION_COOKIE_SECURE=True,
    SESSION_COOKIE_HTTPONLY = True
))


@app.context_processor
def inject_year():
    return dict(year=datetime.datetime.now().strftime("%Y"))


def connect_db():
    """Connects to the specific database."""
    rv = sqlite3.connect(app.config['DATABASE'])
    rv.row_factory = sqlite3.Row
    return rv


def init_db():
    """Initializes the database."""
    with app.open_resource('schema.sql') as f:
        with contextlib.closing(get_db()) as con:
            con.cursor().executescript(f.read())


@app.cli.command('initdb')
def initdb_command():
    """Creates the database tables."""
    init_db()
    print('Initialized the database.')


def get_db():
    """Opens a new database connection if there is none yet for the
    current application context.
    """
    if False:
        if not hasattr(g, 'sqlite_db'):
            g.sqlite_db = connect_db()
        return g.sqlite_db
    else:
        return connect_db()


@app.teardown_appcontext
def close_db(error):
    """Closes the database again at the end of the request."""
    # The database will have been closed by the end of the request, thanks to
    # contextlib.closing().  Besides, the database appears closed at the
    # beginning of the request, probably due to this teardown preserving 
    # the closed connection as g.sqlite_db.
    if False:
        if hasattr(g, 'sqlite_db'):
            g.sqlite_db.close()
            delattr(g, 'sqlite_db')


@app.errorhandler(400)
def handle_bad_request(exc):
    strexc = str(exc)
    log(('Bad request: %s\n%s' % (strexc, traceback.format_exc())).encode("string_escape"), "SUCCESS", "LOW")
    app.logger.exception(strexc)
    return strexc


@app.errorhandler(Exception)
def handle_exception(exc):
    strexc = str(exc)
    log(('Unhandled exception: %s\n%s' % (strexc, traceback.format_exc())).encode("string_escape"), "SUCCESS", "LOW")
    app.logger.exception(strexc)
    return strexc


def get_filepaths(directory):
    """
    This function will generate the file names in a directory
    tree by walking the tree either top-down or bottom-up. For each
    directory in the tree rooted at directory top (including top itself),
    it yields a 3-tuple (dirpath, dirnames, filenames).
    """
    file_paths = []
    for root, directories, files in os.walk(directory):
        for filename in files:
            filepath = os.path.join(root, filename)
            file_paths.append(filepath)
    return file_paths


def get_num(x):
    """get numbers from a string"""
    return int(''.join(ele for ele in x if ele.isdigit()))


def check_version():
    try:
        r = requests.get("http://raw.githubusercontent.com/blabla1337/skf-flask/master/setup.py")
        items_remote = r.content.split(",")
        version_remote = items_remote[1]
        version_remote = version_remote.replace(version_remote[:14], '')
        version_remote = version_remote[:-1]
        with open ("version.txt", "r") as myfile:
            version_local = myfile.read().replace('\n', '')

        if version_local == version_remote:
            return True
        else:
            return False
    except:
        return False


def get_version():
    with open ("version.txt", "r") as myfile:
        version_final = myfile.read().replace('\n', '')
    return version_final


def assert_session():
    if not session.get('logged_in'):
        raise Exception("Not logged in")


def projects_functions_techlist():
    """get list of technology used for creating project functions"""
    assert_session()
    with contextlib.closing(get_db()) as con:
        entries = con.execute('SELECT techID, techName, vulnID from techhacks ORDER BY techID DESC').fetchall()
    return entries


@app.route('/')
@security
def show_landing():
    """show the loging page and set default code language"""
    rand.cleanup()
    csrf_token_raw = rand.bytes(128)
    csrf_token = base64.b64encode(csrf_token_raw)
    session['csrf_token'] = csrf_token
    session['code_lang'] = "php"

    return render_template('login.html', csrf_token=session['csrf_token'])


@app.route('/dashboard', methods=['GET'])
@security
def dashboard():
    """show the landing page"""
    assert_session()
    permissions("read")
    version_check = check_version()
    version = get_version()
    return render_template('dashboard.html', version=version, version_check=version_check)


@app.route('/first-login', methods=['GET'])
@security
def first_login():
    version_check = check_version()
    version = get_version()
    return render_template('first-login.html', version=version, version_check=version_check)


"""create account for a user"""
@app.route('/create-account', methods=['GET', 'POST'])
@security
def create_account():
    """validate the login data for access dashboard page"""
    error = None
    with contextlib.closing(get_db()) as con:
        if request.method == 'POST':
            """Username, password, token, email from form"""
            token  = request.form['token']
            email  = request.form['email']
            password  = request.form['password']
            
            date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")

            #hash the password with Bcrypt, does autosalt
            hashed = bcrypt.generate_password_hash(password,14)

            #check for access
            check = con.execute('SELECT accessToken, userID, activated from users where email=? AND accessToken=?',
                                [email, token]).fetchall()
            for verify in check:
                userID = verify[1]
                if verify[2] == "false":
                    if str(verify[0]) == token:
                        #update the counter and blocker table with new values
                        with con as cur:
                            cur.execute('UPDATE users SET access=?, password=?, activated=? WHERE accessToken=? AND userID=?',
                                           ["true", hashed, "true", token , userID])
                        #Insert record in counter table for the counting of malicious inputs
                        with con as cur:
                            cur.execute('DELETE FROM counter WHERE userID=?',
                                                [userID,])
                            cur.execute('INSERT INTO counter (userID, countEvil, block) VALUES (?, ?, ?)',
                                                [userID, 0, 0])

                        #Create standard group for this user to assign himself to
                        with con as cur:
                            owned_groups = con.execute('SELECT groupName from groups WHERE ownerID=?',
                                    [userID,]).fetchall()
                            owned_group_names = [owned_group[0] for owned_group in owned_groups]
                            if "privateGroup" not in owned_group_names:
                                cur.execute('INSERT INTO groups (ownerID, groupName, timestamp) VALUES (?, ?, ?)',
                                        [userID, "privateGroup", date])

                        #Select this groupID so we can assign the user to this group automatically
                        owned_private_groups = con.execute('SELECT groupID from groups where ownerID=? and groupName=?',
                                                [userID, "privateGroup"]).fetchall()
                        owned_private_groupID = None
                        for owned_private_group in owned_private_groups:
                            owned_private_groupID = owned_private_group[0]

                        #Now we assign the user to the group
                        date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                        with con as cur:
                            cur.execute('INSERT INTO groupMembers (userID, groupID, ownerID) VALUES (?, ?, ?)',
                                                [userID, owned_private_groupID, userID])

        return render_template('login.html', error=error)


"""First comes the method for login"""
@app.route('/login', methods=['GET', 'POST'])
@security
def login():
    """validate the login data for access dashboard page"""
    error = None
    with contextlib.closing(get_db()) as con:
        if request.method == 'POST':
            """Username and password from form"""
            username = request.form['username']
            password = request.form['password']

            #Do DB query also check for access
            check = con.execute('SELECT access from users where userName=?',
                                [username]).fetchall()
            for verify in check:
                if verify[0] == "false":
                    return render_template('warning.html', error=error)

            #Do DB query also check for access
            entries = con.execute('SELECT u.userID, u.privilegeID, u.userName, u.password, u.access, priv.privilegeID, priv.privilege from users as u JOIN privileges AS priv ON priv.privilegeID = u.privilegeID where username=? AND access="true"',
                                [username]).fetchall()
            for entry in entries:
                passwordHash = entry[3]
                userID = entry[0]
                #Do encryption
                if bcrypt.check_password_hash(passwordHash, password):
                    log("Valid username/password submit", "SUCCESS", "HIGH")
                    rand.cleanup()
                    csrf_token_raw = rand.bytes(128)
                    csrf_token = base64.b64encode(csrf_token_raw)
                    session['logged_in'] = True
                    session['userID'] = userID
                    session['csrf_token'] = csrf_token
                    session['code_lang'] = "php"
                    session['userName'] = entry[2]
                    valAlphaNum(session['userName'], 12)
                    session['permissions'] = entry[6]
                    version_check = check_version()
                    version = get_version()

                    #Do DB query also check for access
                    groupID = con.execute('SELECT groupID from groups WHERE groupName=? AND ownerID=?',
                                ["privateGroup", session['userID']]).fetchall()
                    for entry in groupID:
                        session['privateGroup'] = entry[0]
                    return render_template('dashboard.html', version=version, version_check=version_check)
                else:
                    log("invalid login submit", "FAIL", "HIGH")
    return render_template('login.html', error=error)


def countAttempts(counter):
    """We count hacking attempts and block the user if structural"""
    assert_session()

    with contextlib.closing(get_db()) as con:
        entries = con.execute('SELECT * FROM counter where userID=?',
                            [session['userID']]).fetchall()
        counterDB = 0
        blockDB = 0
        for entry in entries:
            counterDB = entry[2]
            blockDB = entry[3]

        updateCount = counterDB + counter
        updateBlock = blockDB + counter
        redirect = False

        if updateCount >= 3:
            countUpdate = 0
            redirect = True

        if updateBlock >= 12:
            redirect = True
            with con as cur:
                cur.execute('UPDATE users SET access=? WHERE userID=?',
                   ["false", session['userID']])
            renderwhat = "/warning.html"

        #update the counter and blocker table with new values
        with con as cur:
            cur.execute('UPDATE counter SET countEvil=?, block=? WHERE userID=?',
                [updateCount, updateBlock, session['userID']])

        if redirect:
            log( "Authenticated session destroyed by counter class", "SUCCESS", "LOW")
            # TO-DO turn on again
            #session.pop('logged_in', None)
            #session.clear()


"""Here is the method for the database enforced privilege based authentication"""
def permissions(fromFunction):
    with contextlib.closing(get_db()) as con:
        """Do DB query to see if username exists"""
        entries = con.execute('SELECT a.username, a.userID, a.password, a.privilegeID, b.privilegeID, b.privilege FROM users as a JOIN privileges as b ON a.privilegeID = b.privilegeID WHERE a.userID =? and a.access="true" ',
                                           [session['userID']]).fetchall()
        perms = ''
        for entry in entries:
            perms = entry[5]

        permissionsGranted = string.split(perms, ':')
        permissionsNeeded  = string.split(fromFunction, ':')

        count = len(permissionsNeeded)
        counthits = 0

        for val in permissionsGranted:
                if val in fromFunction:
                    counthits +=1
        if counthits >= count:
            return perms
        else:
            raise Exception( "User tries to reach functions out of bound no restrictions!!")


@app.route('/logout', methods=['GET', 'POST'])
@security
def logout():
    """logout and destroy session"""
    log( "Authenticated session destroyed", "SUCCESS", "LOW")
    session.pop('logged_in', None)
    session.clear()
    return redirect("/")


@app.route('/code/<code_lang>', methods=['GET'])
@security
def set_code_lang(code_lang):
    """set a code language: php java python perl"""
    assert_session()
    permissions("read")
    allowed = "php java asp"
    valAlphaNum(code_lang, 12)
    found = allowed.find(code_lang)
    if found != -1:
        session['code_lang'] = code_lang
    return redirect(url_for('code_examples'))


@app.route('/code-examples', methods=['GET'])
@security
def code_examples():
    """Shows the knowledge base markdown files."""
    assert_session()
    permissions("read")
    items = []
    id_items = []
    full_file_paths = []
    allowed = set(string.ascii_lowercase + string.ascii_uppercase + '.')
    if set(session['code_lang']) <= allowed:
        full_file_paths = get_filepaths(os.path.join(app.root_path, "markdown/code_examples/"+session['code_lang']))
        for path in full_file_paths:
            basepath = os.path.basename(path)
            baseelems = basepath.split("-")
            id_item = get_num(baseelems[0])
            kb_name_uri = baseelems[-3]
            kb_name = kb_name_uri.replace("_", " ")
            items.append(kb_name)
            id_items.append(id_item)
    return render_template('code-examples.html', items=items, id_items=id_items)


@app.route('/code-item', methods=['POST'])
@security
def show_code_item():
    """show the coding examples page"""
    assert_session()
    permissions("read")
    valNum(request.form['id'], 12)
    id = int(request.form['id'])
    items = []
    full_file_paths = []
    allowed = set(string.ascii_lowercase + string.ascii_uppercase + '.')
    if set(session['code_lang']) <= allowed:
        full_file_paths = get_filepaths(os.path.join(app.root_path, "markdown/code_examples/"+session['code_lang']))
        for path in full_file_paths:
            basepath = os.path.basename(path)
            if id == get_num(basepath.split("-")[0]):
                with open(path, 'r') as codef:
                    codemd = codef.read()
                content = Markup(markdown.markdown(codemd, extensions=['markdown.extensions.fenced_code', 'markdown.extensions.codehilite']))
    return render_template('code-examples-item.html', **locals())


@app.route('/kb-item', methods=['POST'])
@security
def show_kb_item():
    """show the knowledge base search result page"""
    assert_session()
    permissions("read")
    valNum(request.form['id'], 12)
    id = int(request.form['id'])
    kb_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))
    for kbpath in kb_paths:
        basepath = os.path.basename(kbpath)
        if id == get_num(basepath.split("-")[0]):
            with open(kbpath, 'r') as kbpathf:
                kbmd = kbpathf.read()
            content = Markup(markdown.markdown(kbmd))
    return render_template('knowledge-base-item.html', **locals())


@app.route('/knowledge-base-api', methods=['GET'])
@security
def show_kb_api():
    """show the knowledge base items page"""
    log( "User access page /knowledge-base-api", "SUCCESS", "HIGH")
    full_file_paths = []
    content = []
    kb_name = []
    kb_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))
    for kbpath in kb_paths:
        filetmp = open(kbpath, 'r').read()
        filetmp2 = filetmp.replace("-------", "")
        filetmp3 = filetmp2.replace("**", "")
        filetmp4 = filetmp3.replace("\"", "")
        filetmp5 = filetmp4.replace("\t", "")
        content.append(filetmp5.replace("\n", " "))
        basepath = os.path.basename(kbpath)
        kb_name_uri = basepath.split("-")[-3]
        kb_name.append(kb_name_uri.replace("_", " "))
    return render_template('knowledge-base-api.html', **locals())


@app.route('/knowledge-base', methods=['GET'])
@security
def knowledge_base():
    """Shows the knowledge base markdown files."""
    assert_session()
    permissions("read")
    items = []
    id_items = []
    kb_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))
    for kbpath in kb_paths:
        basepath = os.path.basename(kbpath)
        elems = basepath.split("-")
        id_item = get_num(elems[0])
        kb_name_uri = elems[-3]
        kb_name = kb_name_uri.replace("_", " ")
        items.append(kb_name)
        id_items.append(id_item)
    return render_template('knowledge-base.html', items=items, id_items=id_items)


@app.route('/users-new', methods=['GET'])
@security
def user_new():
    """show the create new project page"""
    assert_session()
    permissions("manage")
    return render_template('users-new.html', csrf_token=session['csrf_token'])


@app.route('/users-add', methods=['POST'])
@security
def users_add():
    """add a user"""
    assert_session()
    permissions("manage")
    check_token()
    with contextlib.closing(get_db()) as con:
        with con as cur:
            userName = request.form['username']
            email    = request.form['email']
            privID   = request.form['privID']
            pincode  = request.form['pincode']
            valAlphaNum(userName, 1)
            valNum(privID, 12)
            valNum(pincode, 12)

            users = con.execute('SELECT userID from users WHERE userName=?', [userName,]).fetchall()
            if users:
                raise Exception("User %s already exists" % (userName,))
            users = con.execute('SELECT userID from users WHERE email=?', [email,]).fetchall()
            if users:
                raise Exception("Email %s already registered" % (email,))

            cur.execute('INSERT INTO users (privilegeID, userName, email, password, access, accessToken, activated) VALUES (?, ?, ?, ?, ?, ?, ?)',
                       [privID, userName, email, "none", "false", pincode, "false"])

    return redirect(url_for('users_manage'))


@app.route('/users-manage', methods=['GET'])
@security
def users_manage():
    """show the project list page"""
    assert_session()
    permissions("manage")
    with contextlib.closing(get_db()) as con:
        users = con.execute('SELECT u.userID, u.userName, u.email, u.privilegeID, u.access, p.privilegeID, p.privilege from users as u JOIN privileges as p ON p.privilegeID = u.privilegeID').fetchall()

    return render_template('users-manage.html', users=users, csrf_token=session['csrf_token'])


@app.route('/user-access', methods=['POST'])
@security
def user_access():
    """add a new project to database"""
    assert_session()
    permissions("manage")
    check_token()
    with contextlib.closing(get_db()) as con:
        with con as cur:
            access = request.form['access']
            valBool(access, 12)
            userID = request.form['userID']
            valNum(userID, 12)
            cur.execute('UPDATE users SET access=? WHERE userID=?',
                           [access, userID])
            cur.execute('UPDATE counter SET countEvil=? AND block=? WHERE userID=?',
                           [0, 0, userID])

    return redirect(url_for('users_manage'))


@app.route('/group-new', methods=['GET'])
@security
def group_new():
    """show the create new project page"""
    assert_session()
    permissions("edit")
    return render_template('group-new.html', csrf_token=session['csrf_token'])


@app.route('/group-add', methods=['POST'])
@security
def group_add():
    """create a group"""
    assert_session()
    permissions("edit")
    check_token()
    with contextlib.closing(get_db()) as con:
        inputName = request.form['groupName']
        valAlphaNum(inputName, 3)
        date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")

        groups = con.execute('SELECT groupID from groups WHERE groupName=?',
                            [inputName,]).fetchall()
        for group in groups:
            raise Exception("Group %s already exists" % (inputName,))

        with con as cur:
            cur.execute('INSERT INTO groups (timestamp, groupName, ownerID) VALUES (?, ?, ?)',
                   [date, inputName, session['userID']])
        # select the latest group in order to check id
        groups = con.execute('SELECT groupID from groups WHERE timestamp=? AND groupName=? AND ownerID=?',
                            [date, inputName, session['userID']]).fetchall()
        groupID = None
        for group in groups:
            groupID = group[0]

        # insert this back into groupMembers table so the user is added to group
        with con as cur:
            cur.execute('INSERT INTO groupMembers (userID, groupID, ownerID, timestamp) VALUES (?, ?, ?, ?)',
                   [session['userID'], groupID, session['userID'], date])
    return redirect(url_for('group_manage'))


@app.route('/group-users', methods=['GET'])
@security
def group_users():
    """show the project list page"""
    assert_session()
    permissions("edit")
    with contextlib.closing(get_db()) as con:
        groups = con.execute('SELECT * from groups where ownerID=? and groupName !=? ',
                              [session['userID'], "privateGroup"]).fetchall()

        """Select all users for adding to group"""
        users = con.execute('SELECT username, userID from users').fetchall()

        """select users by assigned groups for display"""
        summary = con.execute('SELECT u.username, u.userID, g.groupName, g.groupID, m.groupID, m.userID, m.timestamp, g.ownerID from users as u JOIN groups AS g ON g.groupID = m.groupID JOIN groupMembers as m ON u.userID = m.userID  WHERE g.ownerID=? AND u.userName !=? ORDER BY g.groupName ',
                                       [session['userID'], session['userName']]).fetchall()

    return render_template('group-users.html', groups=groups, users=users, summary=summary, csrf_token=session['csrf_token'])


@app.route('/group-add-users', methods=['POST'])
@security
def group_add_users():
    """add a user to the group"""
    assert_session()
    permissions("edit")
    check_token()
    f = request.form
    groupID = f['groupID']
    valNum(groupID, 12)
    date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")

    """Check if submitted groupID is owned by user"""
    with contextlib.closing(get_db()) as con:
        owngroups = con.execute('SELECT groupID from groups where ownerID=?',
                                       [session['userID']]).fetchall()
        for owngroup in owngroups:
            if int(groupID) == int(owngroup[0]):
                members = con.execute('SELECT userID from groupMembers WHERE groupID=?',
                                       [groupID,]).fetchall()
                memberIDs = [str(member[0]) for member in members]
                for key in f.keys():
                    if key.find("test") >= 0:
                        for value in f.getlist(key):
                            items = value.split("-")
                            userID = items[0]
                            valNum(userID, 12)
                            if userID in memberIDs:
                                print "User ID %s already belongs in groupID %s" % (userID, groupID)
                            else:
                                with con as cur:
                                    cur.execute('INSERT INTO groupMembers (timestamp, groupID, userID, ownerID) VALUES (?, ?, ?, ?)',
                                           [date, groupID, userID, session['userID']])
                                memberIDs.append(userID)

    redirect_url = '/group-users'
    return redirect(redirect_url)


@app.route('/group-del-users', methods=['POST'])
@security
def group_del_users():
    """add a user from the group"""
    assert_session()
    permissions("edit")
    check_token()
    f = request.form
    
    userID = f['userID']
    valNum(userID, 12)

    groupID = f['groupID']
    valNum(groupID, 12)

    """Check if submitted groupID is owned by user"""
    with contextlib.closing(get_db()) as con:
        owngroups = con.execute('SELECT groupID from groups where ownerID=?',
                                       [session['userID']]).fetchall()
        for owngroup in owngroups:
            if int(groupID) == int(owngroup[0]):
                members = con.execute('SELECT userID from groupMembers WHERE groupID=?',
                                       [groupID,]).fetchall()
                memberIDs = [str(member[0]) for member in members]
                if userID not in memberIDs:
                    print "User ID %s does not belong in groupID %s" % (userID, groupID)
                else:
                    with con as cur:
                        cur.execute('DELETE FROM groupMembers WHERE groupID=? and userID=?',
                           [groupID, userID])
                    memberIDs = [memberID for memberID in memberIDs if memberID != userID]

    redirect_url = '/group-users'
    return redirect(redirect_url)


@app.route('/user-del', methods=['POST'])
@security
def user_del():
    """delete user"""
    assert_session()
    permissions("delete")
    check_token()
    userID  = request.form['userID']
    valNum(userID, 12)

    with contextlib.closing(get_db()) as con:
        with con as cur:
            cur.execute("DELETE FROM users WHERE userID=?",
                [userID])
            cur.execute("DELETE FROM groupMembers WHERE userID=?",
                [userID])
            cur.execute("DELETE FROM groups WHERE ownerID=?",
                [userID])
    return redirect("/users-manage")


@app.route('/group-manage', methods=['GET'])
@security
def group_manage():
    """show the project list page"""
    assert_session()
    permissions("edit")
    with contextlib.closing(get_db()) as con:
        groups = con.execute('SELECT * from groups where ownerID=? and groupName !=? ',
                          [session['userID'], "privateGroup"]).fetchall()

    return render_template('group-manage.html', groups=groups, csrf_token=session['csrf_token'])


@app.route('/group-del', methods=['POST'])
@security
def group_del():
    """delete project from database"""
    assert_session()
    permissions("manage")
    check_token()
    groupID = request.form['groupID']
    valNum(groupID, 12)

    with contextlib.closing(get_db()) as con:
        with con as cur:
            cur.execute("DELETE FROM groups WHERE groupID=? AND ownerID=?",
                [groupID, session['userID']])
    return redirect("/group-manage")


@app.route('/project-new', methods=['GET'])
@security
def projects():
    """show the create new project page"""
    assert_session()
    permissions("edit")
    return render_template('project-new.html', csrf_token=session['csrf_token'])


@app.route('/project-add', methods=['POST'])
@security
def add_entry():
    """add a new project to database"""
    assert_session()
    permissions("edit")
    check_token()
    with contextlib.closing(get_db()) as con:
        inputName = request.form['inputName']
        inputVersion = request.form['inputVersion']
        inputDesc = request.form['inputDesc']
        # valAlphaNum(inputName, 1)
        # valNum(inputVersion, 1)
        date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
        with con as cur:
            cur.execute('INSERT INTO projects (timestamp, projectName, projectVersion, projectDesc, userID, ownerID, groupID) VALUES (?, ?, ?, ?, ?, ?, ?)',
                   [date, inputName, inputVersion, inputDesc, session['userID'],  session['userID'], session['privateGroup']])
    return redirect(url_for('project_list'))


@app.route('/assign-group', methods=['POST'])
@security
def assign_group():
    """add a new project to database"""
    assert_session()
    permissions("edit")
    check_token()
    with contextlib.closing(get_db()) as con:
        groupID   = request.form['groupID']
        projectID = request.form['projectID']
        valNum(projectID, 12)
        valNum(groupID, 12)
        """Check is submitted groupID is owned by user"""
        owner = con.execute('SELECT groupID from groups where ownerID=?',
                                       [session['userID']]).fetchall()
        for val in owner:
            if int(groupID) == int(val[0]):
                with con as cur:
                    cur.execute('UPDATE projects SET groupID=? WHERE projectID=? AND userID=?',
                       [groupID, projectID, session['userID']])
    return redirect(url_for('project_list'))


@app.route('/project-del', methods=['POST'])
@security
def project_del():
    """delete project from database"""
    assert_session()
    permissions("delete")
    id = request.form['projectID']
    valNum(id, 12)
    check_token()
    with contextlib.closing(get_db()) as con:
        with con as cur:
            cur.execute("DELETE FROM projects WHERE projectID=? AND userID=? AND ownerID=?",
               [id, session['userID'], session['userID']])
    return redirect("/project-list")


@app.route('/project-list', methods=['GET'])
@security
def project_list():
    """show the project list page"""
    assert_session()
    permissions("read")
    with contextlib.closing(get_db()) as con:
        #First query is for the users own owned projects
        entries = con.execute('SELECT p.projectName, p.projectVersion, p.projectDESC, p.projectID, p.timestamp, p.groupID, g.groupName, g.groupID FROM projects as p JOIN groups as g ON g.groupID = p.groupID where p.userID=? ORDER BY projectID DESC',
                              [session['userID']]).fetchall()
        #select the groups which can be selected by this user
        groups = con.execute('SELECT * FROM groups WHERE ownerID=?',
                              [session['userID']]).fetchall()
    return render_template('project-list.html', entries=entries, groups=groups, csrf_token=session['csrf_token'])


@app.route('/project-shared', methods=['GET'])
@security
def project_shared():
    """show the project list page"""
    assert_session()
    permissions("read")
    with contextlib.closing(get_db()) as con:
        #Here we see what projects this users was assigned to
        entries = con.execute('SELECT p.projectName, p.projectVersion, p.projectDESC, p.projectID, p.timestamp, p.groupID, p.ownerID, m.userID, m.groupID, u.userID, u.userName FROM projects as p JOIN groupMembers as m ON m.groupID = p.groupID JOIN users as u ON u.userID=p.ownerID where m.userID=? AND u.userName !=? ORDER BY p.projectID DESC',
                              [session['userID'], session['userName']]).fetchall()

    return render_template('project-shared.html', entries=entries, csrf_token=session['csrf_token'])


@app.route('/project-options/<project_id>', methods=['GET'])
@security
def projects_options(project_id):
    """show the project options landing page"""
    assert_session()
    permissions("read")
    valNum(project_id, 12)
    return render_template('project-options.html', project_id=project_id, csrf_token=session['csrf_token'])


@app.route('/project-functions/<project_id>', methods=['GET'])
@security
def project_functions(project_id):
    """show the pproject functions page"""
    assert_session()
    permissions("read")
    techlist = projects_functions_techlist()
    valNum(project_id, 12)
    with contextlib.closing(get_db()) as con:
        entries = con.execute('SELECT p.paramID, p.functionName, p.functionDesc, p.projectID, p.userID, p.tech, p.techVuln, p.entryDate, t.techName, proj.projectID, proj.groupID, m.userID, m.groupID FROM parameters AS p JOIN techhacks AS t ON p.tech = t.techID JOIN projects as proj ON proj.projectID = p.projectID JOIN groupMembers as m ON m.groupID = proj.groupID WHERE proj.projectID=? AND m.userID=? GROUP BY t.techName',
                      [project_id, session['userID']]).fetchall()
    return render_template('project-functions.html', project_id=project_id, techlist=projects_functions_techlist(), entries=entries, csrf_token=session['csrf_token'])


@app.route('/project-function-del', methods=['POST'])
@security
def function_del():
    """delete a project function"""
    assert_session()
    permissions("delete")
    check_token()
    projectID = request.form['projectID']
    paramID = request.form['paramID']
    valNum(projectID, 12)
    valNum(paramID, 12)
    with contextlib.closing(get_db()) as con:
        #First check if the user is allowed to delete this parameter
        projects = con.execute('SELECT p.projectID, p.groupID, m.groupID, m.userID from projects as p '
            'JOIN groupMembers as m ON m.groupID = p.groupID where m.userID=?',
            [session['userID']]).fetchall()
        for project in projects:
            if int(projectID) == int(project[0]):
                with con as cur:
                    cur.execute("DELETE FROM parameters WHERE projectID=? AND paramID=?",
                                   [projectID, paramID])
                break
    redirect_url = "/project-functions/" + projectID
    return redirect(redirect_url)


@app.route('/project-function-add', methods=['POST'])
@security
def add_function():
    """add a project function"""
    assert_session()
    permissions("edit")
    check_token()
    id = request.form['project_id']
    valNum(id, 12)
    fName = request.form['functionName']
    # valAlphaNum(fName, 1)
    fDesc = request.form['functionDesc']
    # valAlphaNum(fDesc, 1)

    #Check is submitted projectID is owned by user
    with contextlib.closing(get_db()) as con:
        projects = con.execute('SELECT p.projectID, p.groupID, m.groupID, m.userID from projects as p JOIN groupMembers as m ON m.groupID = p.groupID where m.userID=?',
                                       [session['userID']]).fetchall()
        for project in projects:
            if int(id) == int(project[0]):
                f = request.form
                for key in f.keys():
                    for value in f.getlist(key):
                        found = key.find("test")
                        if found != -1:
                            date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                            items = value.split("-")
                            techID = items[2]
                            vulnID = items[0]
                            valAlphaNum(techID, 12)
                            valAlphaNum(vulnID, 12)
                            with con as cur:
                                cur.execute('INSERT INTO parameters (entryDate, functionName, functionDesc, techVuln, tech, projectID, userID) VALUES (?, ?, ?, ?, ?, ?, ?)',
                                       [date, fName, fDesc, vulnID, techID, id, session['userID']])
    redirect_url = '/project-functions/'+str(id)
    return redirect(redirect_url)


CKELEM_PREFIX = "vulnID"
CKELEM_PREFIX_LEN = len(CKELEM_PREFIX)

@app.route('/project-checklist-add', methods=['POST'])
@security
def add_checklist():
    """add project checklist"""
    assert_session()
    permissions("edit")

    # https://github.com/pallets/werkzeug/issues/1068
    request.get_data()
    
    # log("Type of request: %s" % (type(request),), "SUCCESS", "LOW")
    # log("dir(request): %s" % (dir(request),), "SUCCESS", "LOW")
    # log("Type of form: %s" % (type(request.form),), "SUCCESS", "LOW")
    # for key in request.form:
    #     log("Form element: \"%s\"=\"%s\"" % (key, request.form[key],), "SUCCESS", "LOW")
    # log("Request form csrf token: %s" % (request.form["csrf_token"],), "SUCCESS", "LOW")
    check_token()
    #We do valNum for projectID here because we need it in the comparison
    project_id = request.form['projectID']
    valNum(project_id, 12)
    project_name = request.form['projectName']
    valAlphaNum(project_name, 12)

    with contextlib.closing(get_db()) as con:
        #Check if submitted projectID is owned by user
        projects = con.execute('SELECT p.projectID, p.groupID, m.groupID, m.userID from projects as p JOIN groupMembers as m ON m.groupID = p.groupID where m.userID=?',
    				   [session['userID']]).fetchall()
        for project in projects:
            if int(project_id) == int(project[0]):
                f = request.form
                qnumbers = []
                for key in f.keys():
                    if key.startswith(CKELEM_PREFIX):
                        for value in f.getlist(key):
                            qnostr = key[CKELEM_PREFIX_LEN:]
                            valNum(qnostr, 12)
                            qnumbers.append(int(qnostr))
                qnumbers.sort()
                for qno in qnumbers:
                        qnostr = str(qno)

                        listidx = "listID" + qnostr
                        # The questionlist table has a listID field which is just an incrementing index.
                        # This form value carries the list name to be placed into the listName field.
                        list_name = f[listidx]
                        valAlphaNum(list_name, 12)
                        
                        answeridx = "answer" + qnostr
                        answer = f[answeridx]
                        valAlphaNum(answer, 12)
                        
                        questionidx = "questionID" + qnostr
                        questionID = f[questionidx]
                        valNum(questionID, 12)
                        
                        vulnidx = "vulnID" + qnostr
                        vulnID = f[vulnidx]
                        valNum(vulnID, 12)
                        
                        date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                        with con as cur:
                            cur.execute('INSERT INTO questionlist (entryDate, answer, projectName, projectID, questionID, vulnID, listName, userID) VALUES (?, ?, ?, ?, ?, ?, ?, ?)',
                                    [date, answer, project_name, project_id, questionID, vulnID, 
                                        list_name, session['userID']])
    redirect_url = "/results-checklists"
    return redirect(redirect_url)


def populate_checklists(checklist_paths, kb_paths, lvl_name, 
        owasp_ids_lvl, owasp_ygbs_lvl, owasp_kbs_lvl, owasp_contents_lvl, owasp_descs_lvl):
    owasp_ids = []
    owasp_ygbs = []
    owasp_kbs = []
    owasp_contents = []
    owasp_descs = []

    for path in checklist_paths:
        basepath = os.path.basename(path)
        owasp_path_elems = basepath.split("-")
        owasp_checklist_name = owasp_path_elems[2]
        if owasp_checklist_name == "ASVS":
            owasp_checklist_name = "-".join(owasp_path_elems[2:5])
            parse_index = 6
        else:
            parse_index = 4
        if owasp_checklist_name == lvl_name:
            owasp_id = get_num(owasp_path_elems[0])
            owasp_kb = owasp_path_elems[parse_index]
            owasp_ygb = owasp_path_elems[parse_index + 2]

            owasp_ids.append(owasp_id)
            owasp_kbs.append(owasp_kb)
            owasp_ygbs.append(owasp_ygb)
            with open(path, 'r') as pathf:
                checklistmd = pathf.read()
            owasp_contents.append(Markup(markdown.markdown(checklistmd)))

            descriptions = []
            for kbpath in kb_paths:
                kbbasepath = os.path.basename(kbpath)
                path_vuln = get_num(kbbasepath.split("-")[0])
                if int(owasp_kb) == int(path_vuln):
                    with open(kbpath, 'r') as kbpathf:
                        kbmd = kbpathf.read()
                    description = kbmd.split("**")
                    descriptions.append(description[2])
            owasp_descs.append("\n".join(descriptions))

    owasp_ids_lvl.append(owasp_ids)
    owasp_ygbs_lvl.append(owasp_ygbs)
    owasp_kbs_lvl.append(owasp_kbs)
    owasp_contents_lvl.append(owasp_contents)
    owasp_descs_lvl.append(owasp_descs)

    return lvl_name
    

NUM_ASVS_LEVELS = 6
ASVS_1, ASVS_2, ASVS_3, ASVS_BASIC, ASVS_ADVANCED, ASVS_CUSTOM = range(NUM_ASVS_LEVELS)
ASVS_NAMES = ("ASVS-level-1", "ASVS-level-2", "ASVS-level-3", "CS_basic_audit", "CS_advanced_audit", "custom")
ASVS_TITLES = ("OWASP ASVS Level 1", "OWASP ASVS Level 2", "OWASP ASVS Level 3", 
        "Basic Audit Checklist", "Advanced Audit Checklist", "Custom Checklist")


@app.route('/project-checklists/<project_id>', methods=['GET'])
@security
def project_checklists(project_id):
    """show the project checklists page"""
    assert_session()
    permissions("read")
    valNum(project_id, 12)
    with contextlib.closing(get_db()) as con:
        projects = con.execute('SELECT p.projectID, p.userID, p.groupID, p.projectName, p.projectVersion, p.projectDesc, p.ownerID, m.userID, m.groupID FROM projects as p JOIN groupMembers AS m ON m.groupID = p.groupID WHERE p.projectID=? AND m.userID=?',
                            [project_id, session['userID']]).fetchall()
    projectName = ""

    owasp_level_descs = []
    owasp_level_recommendations = []

    owasp_ids_lvl = []
    owasp_ygbs_lvl = []
    owasp_kbs_lvl = []
    owasp_contents_lvl = []
    owasp_descs_lvl = []

    for prep in projects:
        projectName = prep[3]
        
        checklist_paths = get_filepaths(os.path.join(app.root_path, "markdown/checklists"))
        checklist_paths.sort()

        kb_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))
        kb_paths.sort()

        for level0 in range(NUM_ASVS_LEVELS):
            level = level0 + 1
            level_name = ASVS_NAMES[level0]
            populate_checklists(checklist_paths, kb_paths, level_name,
                    owasp_ids_lvl, owasp_ygbs_lvl, 
                    owasp_kbs_lvl, owasp_contents_lvl, owasp_descs_lvl)
            if level0 < 3:
                level_desc = "OWASP Application Security Verification Standard Level %d" % (level0 + 1,)
                if level0 < 2:
                    level_recommendation = "Recommended"
                else:
                    level_recommendation = "Advanced"
            else:
                level_desc = ("This checklist is a template for your own %s checklist. "
                        "If you have created one please create a Pull request on GIT.") % (level_name,)
                level_recommendation = "Custom"
            owasp_level_descs.append(level_desc)
            owasp_level_recommendations.append(level_recommendation)

        break

    return render_template('project-checklists.html', csrf_token=session['csrf_token'], 
                NUM_ASVS_LEVELS=NUM_ASVS_LEVELS,
                ASVS_NAMES=ASVS_NAMES,
                ASVS_TITLES=ASVS_TITLES,
                owasp_level_descs=owasp_level_descs,
                owasp_level_recommendations=owasp_level_recommendations,
                owasp_ids_lvl=owasp_ids_lvl,
                owasp_ygbs_lvl=owasp_ygbs_lvl, 
                owasp_kbs_lvl=owasp_kbs_lvl, 
                owasp_contents_lvl=owasp_contents_lvl, 
                owasp_descs_lvl=owasp_descs_lvl,
                projectName=projectName,
                project_id=project_id
            )


@app.route('/results-checklists', methods=['GET'])
@security
def results_checklists():
    """show the results checklists page"""
    assert_session()
    permissions("read")
    with contextlib.closing(get_db()) as con:
        entries = con.execute('SELECT p.projectID, p.projectName, p.projectVersion, q.listName, q.entryDate FROM questionlist AS q JOIN projects AS p ON q.projectID = p.projectID JOIN groupMembers as m ON m.groupID = p.groupID WHERE m.userID=? GROUP BY p.projectName, p.projectVersion, q.listName, q.entryDate ORDER BY p.projectName ASC, p.projectVersion DESC, q.listName ASC, q.entryDate DESC',
                          [session['userID']]).fetchall()
    return render_template('results-checklists.html', entries=entries, csrf_token=session['csrf_token'])


@app.route('/results-functions', methods=['GET'])
@security
def results_functions():
    """show the results functions page"""
    assert_session()
    permissions("read")
    with contextlib.closing(get_db()) as con:
        entries = con.execute('SELECT p.projectID, p.projectName, p.projectVersion, par.functionName, par.entryDate FROM projects AS p JOIN parameters AS par on p.projectID = par.projectID JOIN groupMembers AS m ON m.groupID = p.groupID WHERE m.userID=? GROUP BY p.projectName, p.projectVersion, par.functionName, par.entryDate ORDER BY p.projectName ASC, p.projectVersion DESC, par.entryDate DESC, par.functionName ASC',
                         [session['userID']]).fetchall()
    return render_template('results-functions.html', entries=entries, csrf_token=session['csrf_token'])


@app.route('/results-functions-del', methods=['POST'])
@security
def functions_del():
    """delete functions result items"""
    assert_session()
    permissions("delete")
    check_token()
    entryDate = request.form['entryDate']
    projectID = request.form['projectID']
    valNum(projectID, 12)
    with contextlib.closing(get_db()) as con:
        #Use select in order to see if this user is linked to project
        entries = con.execute("SELECT p.projectID, p.groupID, m.groupID, m.userID FROM projects AS p JOIN groupMembers AS m ON m.groupID = p.groupID WHERE m.userID=?  ",
                            [session['userID']]).fetchall()
        for entry in entries:
            if int(entry[0]) == int(projectID):
                with con as cur:
                    cur.execute("DELETE FROM parameters WHERE entryDate=? AND projectID=?",
                        [entryDate, projectID])
    return redirect("/results-functions")


@app.route('/results-checklists-del', methods=['POST'])
@security
def checklists_del():
    """delete checklist result item"""
    assert_session()
    permissions("delete")
    check_token()
    entryDate = request.form['entryDate']
    projectID = request.form['projectID']
    valNum(projectID, 12)
    with contextlib.closing(get_db()) as con:
        #Use select in order to see if this user is linked to project
        entries = con.execute("SELECT p.projectID, p.groupID, m.groupID, m.userID FROM projects AS p JOIN groupMembers AS m ON m.groupID = p.groupID WHERE m.userID=?  ",
                            [session['userID']]).fetchall()
        for entry in entries:
            if int(entry[0]) == int(projectID):
                with con as cur:
                    cur.execute("DELETE FROM questionlist WHERE entryDate=? AND projectID=? ",
                           [entryDate, projectID])
    return redirect("/results-checklists")


@app.route('/results-checklist-report/<entryDate>', methods=['GET'])
@security
def checklist_results(entryDate):
    """show checklist results report"""
    assert_session()
    permissions("read")
    ygb = []
    questions = []
    content = []
    with contextlib.closing(get_db()) as con:
        entries = con.execute("SELECT l.projectID, p.projectName, p.projectVersion, l.listName, l.questionID, l.vulnID FROM questionlist AS l JOIN projects AS p ON p.projectID = l.projectID JOIN groupMembers AS m ON m.groupID = p.groupID WHERE l.answer='no' AND l.entryDate=? AND m.userID=? ORDER BY p.projectName, p.projectVersion, l.listName, l.questionID",
               [entryDate, session['userID']]).fetchall()
    projectName = None
    projectVersion = None
    listName = None
    questionID = None
    vulnID = None
    for entry in entries:
        projectName = entry[1]
        projectVersion = entry[2]
        listName = entry[3]
        questionID = entry[4]
        vulnID = entry[5]
        kb_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))
        for kbpath in kb_paths:
            kbbasepath = os.path.basename(kbpath)
            kbpath_vuln = get_num(kbbasepath.split("-")[0])
            if int(vulnID) == int(kbpath_vuln):
                with open(kbpath, 'r') as kbpathf:
                    kbmd = kbpathf.read()
                content.append(Markup(markdown.markdown(kbmd)))

                checklist_paths = get_filepaths(os.path.join(app.root_path, "markdown/checklists"))
                ckqs = []
                ckygbs = []
                for path in checklist_paths:
                    basepath = os.path.basename(path)
                    elems = basepath.split("-")
                    path_questionID = get_num(elems[0])
                    if int(questionID) == int(path_questionID):
                        with open(path, 'r') as pathf:
                            checklistmd = pathf.read()
                        ckqs.append(Markup(markdown.markdown(checklistmd)))
                        checklist_name = elems[2]
                        if checklist_name == "ASVS":
                            checklist_name = "-".join(elems[2:5])
                            checklist_kb = elems[6]
                            checklist_ygb = elems[8]
                        else:
                            checklist_kb = elems[4]
                            checklist_ygb = elems[6]
                        ckygbs.append(checklist_ygb)
                questions.append("\n".join(ckqs))
                ygb.append(" ".join(ckygbs))

    return render_template('results-checklist-report.html', **locals())


@app.route('/results-checklist-docx/<entryDate>')
def download_file_checklist(entryDate):
    """Download checklist results report in docx"""
    assert_session()
    permissions("read")
    ygb_docx = []
    content_raw = []
    content_checklist = []
    content_title = []
    with contextlib.closing(get_db()) as con:
        entries = con.execute("SELECT l.projectID, p.projectName, p.projectVersion, l.listName, l.questionID, l.vulnID FROM questionlist AS l JOIN projects AS p ON p.projectID = l.projectID JOIN groupMembers AS m ON m.groupID = p.groupID WHERE l.answer='no' AND l.entryDate=? AND m.userID=? ORDER BY p.projectName, p.projectVersion, l.listName, l.questionID",
               [entryDate, session['userID']]).fetchall()
    document = Document()
    document.add_picture(os.path.join(app.root_path,'static/img/banner-docx.jpg'), width=Inches(5.125), height=Inches(1.042))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    #document.add_heading('Security Knowledge Framework', 0)
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph()
    projectName = None
    projectVersion = None
    listName = None
    questionID = None
    vulnID = None
    for entry in entries:
        projectName = entry[1]
        projectVersion = entry[2]
        listName = entry[3]
    p.add_run('Project '  + str(projectName) + ' ' + str(projectVersion))
    p.add_run('\r\n')
    p.add_run('Checklist ' + str(listName) + ' answers as of ' + entryDate)
    document.add_page_break()
    p = document.add_heading('Table of contents', level=1)
    p.add_run('\r\n')
    document.add_paragraph('Introduction')
    questionID = None
    vulnID = None

    for entry in entries:
        questionID = entry[4]
        vulnID = entry[5]
        kb_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))
        for kbpath in kb_paths:
            kbbasepath = os.path.basename(kbpath)
            kbpath_vuln = get_num(kbbasepath.split("-")[0])
            if int(vulnID) == int(kbpath_vuln):
                with open(kbpath, 'r') as kbpathf:
                    kbmd = kbpathf.read()
                content = Markup(markdown.markdown(kbmd))
                text = ''.join(BeautifulSoup(content).findAll(text=True))
                text_encode = text.encode('utf-8')
                content_title.append(text_encode.splitlines()[0])
                text_encode = text_encode.replace("Solution", "\nSolution");
                content_raw.append(text_encode)

                checklist_paths = get_filepaths(os.path.join(app.root_path, "markdown/checklists"))
                ckqs = []
                ckygbs = []
                for path in checklist_paths:
                    basepath = os.path.basename(path)
                    elems = basepath.split("-")
                    path_questionID = get_num(elems[0])
                    if int(questionID) == int(path_questionID):
                        with open(path, 'r') as pathf:
                            checklistmd = pathf.read()
                        ckqs.append(Markup(markdown.markdown(checklistmd)))
                        checklist_name = elems[2]
                        if checklist_name == "ASVS":
                            checklist_name = "-".join(elems[2:5])
                            checklist_kb = elems[6]
                            checklist_ygb = elems[8]
                        else:
                            checklist_kb = elems[4]
                            checklist_ygb = elems[6]
                        ckygbs.append(checklist_ygb)
                content_checklist.append("\n".join(ckqs))
                ygb_docx.append(" ".join(ckygbs))

    for item in content_title:
        p = document.add_paragraph(item)
        p.add_run()
    document.add_page_break()
    document.add_heading('Introduction', level=1)
    p = document.add_paragraph(
        'The security knowledge framework is composed by means of the highest security standards currently available and is designed to maintain the integrity of your application, so you and your costumers sensitive data is protected against hackers. This document is provided with a checklist in which the programmers of your application had to run through in order to provide a secure product.'
    )
    p.add_run('\n')
    p = document.add_paragraph(
        'In the post-development stage of the security knowledge framework the developer double-checks his application against a checklist which consists out of several questions asking the developer about different stages of development and the methodology of implementing different types of functionality the application contains. After filling in this checklist the developer gains feedback on the failed checklist items providing him with solutions about how to solve the additional vulnerability\'s found in the application.'
    )
    document.add_page_break()
    i = 0
    for item in content_raw:
        document.add_heading(content_title[i], level=1)
        result = re.sub("<p>", " ", content_checklist[i])
        result1 = re.sub("</p>", " ", result)
        document.add_heading(result1, level=4)
        p = document.add_paragraph(item.partition("\n")[2])
        ygb = ygb_docx[i]
        if "y" in ygb:
            document.add_picture(os.path.join(app.root_path,'static/img/yellow.png'), width=Inches(0.20))
        if "g" in ygb:
            document.add_picture(os.path.join(app.root_path,'static/img/green.png'), width=Inches(0.20))
        if "b" in ygb:
            document.add_picture(os.path.join(app.root_path,'static/img/blue.png'), width=Inches(0.20))
        p.add_run("\n")
        document.add_page_break()
        i += 1
    # FIXME: save in a request-specific location or in memory
    document.save("checklist-security-report.docx")
    headers = {"Content-Disposition": "attachment; filename=%s" % "checklist-security-report.docx"}
    file_path = os.path.join(app.root_path, "checklist-security-report.docx")
    with open("checklist-security-report.docx", 'rb') as f:
        body = f.read()
    return make_response((body, headers))


@app.route('/results-function-report/<projectID>', methods=['GET'])
@security
def function_results(projectID):
    """show checklist results report"""
    assert_session()
    permissions("read")
    content = []
    valNum(projectID, 12)
    with contextlib.closing(get_db()) as con:
        entries = con.execute("SELECT projects.projectName, projects.projectID, projects.projectVersion, parameters.functionName, parameters.tech, parameters.functionDesc, parameters.entryDate, parameters.techVuln, techhacks.techName, projects.userID, projects.groupID, m.userID, m.groupID FROM projects JOIN parameters ON parameters.projectID=projects.projectID JOIN techhacks ON techhacks.techID = parameters.tech JOIN groupMembers AS m ON m.groupID = projects.groupID WHERE parameters.projectID=? AND m.userID=? GROUP BY parameters.tech ORDER BY parameters.tech ASC",
               [projectID, session['userID']]).fetchall()
    projectName = None
    projectVersion = None
    functionName = None
    functionDesc = None
    vulnID = None
    for entry in entries:
        projectName = entry[0]
        projectVersion = entry[2]
        functionName = entry[3]
        functionDesc = entry[5]
        vulnID = entry[7]
        kb_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))
        for kbpath in kb_paths:
            kbbasepath = os.path.basename(kbpath)
            kbpath_vuln = get_num(kbbasepath.split("-")[0])
            if int(vulnID) == int(kbpath_vuln):
                with open(kbpath, 'r') as kbpathf:
                    kbmd = kbpathf.read()
                content.append(Markup(markdown.markdown(kbmd)))
    return render_template('results-function-report.html', **locals())


@app.route('/results-function-docx/<projectID>')
def download_file_function(projectID):
    """Download checklist results report in docx"""
    assert_session()
    permissions("read")
    content_raw = []
    content_title = []
    content_tech = []
    valNum(projectID, 12)
    with contextlib.closing(get_db()) as con:
        entries = con.execute("SELECT projects.projectName, projects.projectID, projects.projectVersion, parameters.functionName, parameters.tech, parameters.functionDesc, parameters.entryDate, parameters.techVuln, techhacks.techName, projects.userID, projects.groupID, m.userID, m.groupID FROM projects JOIN parameters ON parameters.projectID=projects.projectID JOIN techhacks ON techhacks.techID = parameters.tech JOIN groupMembers AS m ON m.groupID = projects.groupID WHERE parameters.projectID=? AND m.userID=? GROUP BY parameters.tech ORDER BY parameters.tech ASC",
               [projectID, session['userID']]).fetchall()
    document = Document()
    document.add_picture(os.path.join(app.root_path,'static/img/banner-docx.jpg'), width=Inches(5.125), height=Inches(1.042))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    #document.add_heading('Security Knowledge Framework', 0)
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph()
    projectName = None
    projectVersion = None
    functionName = None
    functionDesc = None
    for entry in entries:
        projectName = entry[0]
        projectVersion = entry[2]
        functionName = entry[3]
        functionDesc = entry[5]
    p.add_run('Date: ' + datetime.datetime.now().strftime("%Y-%m-%d %H:%M"))
    p.add_run('\r\n')
    p.add_run('Project: ' + str(projectName) + ' ' + str(projectVersion))
    p.add_run('\r\n')
    p.add_run('Function: ' + str(functionName))
    p.add_run('\r\n')
    p.add_run(str(functionDesc))
    document.add_page_break()
    p = document.add_heading('Table of contents', level=1)
    p.add_run('\r\n')
    document.add_paragraph('Introduction')
    vulnID = None
    for entry in entries:
        vulnID = entry[7]
        kb_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))
        for kbpath in kb_paths:
            kbbasepath = os.path.basename(kbpath)
            kbpath_vuln = get_num(kbbasepath.split("-")[0])
            if int(vulnID) == int(kbpath_vuln):
                with open(kbpath, 'r') as kbpathf:
                    kbmd = kbpathf.read()
                content = Markup(markdown.markdown(kbmd))
                text = ''.join(BeautifulSoup(content).findAll(text=True))
                text_encode = text.encode('utf-8')
                content_title.append(text_encode.splitlines()[0])
                text_encode = text_encode.replace("Solution", "\nSolution");
                content_raw.append(text_encode)
    for item in content_title:
        p = document.add_paragraph(item)
        p.add_run()
    document.add_page_break()
    document.add_heading('Introduction', level=1)
    p = document.add_paragraph(
        'The security knowledge framework is composed by means of the highest security standards currently available and is designed to maintain the integrity of your application, so you and your costumers sensitive data is protected against hackers. This document is provided with a checklist in which the programmers of your application had to run through in order to provide a secure product.'
    )
    p.add_run('\n')
    p = document.add_paragraph(
        'In this part of security knowledge framework, all the parameters and variables are audited by means of the information given by the programmer such as the processing techniques. Each of these techniques contain different types of vulnerabilities when implemented in a improper fashion. This document will raise awareness about these vulnerabilities, as well as presenting solutions for the right implementation.'
    )
    document.add_page_break()
    i = 0
    for item in content_raw:
        document.add_heading("Knowledge-Base: "+content_title[i], level=1)
        document.add_heading("Technology: "+entries[i][8], level=2)
        p = document.add_paragraph(item.partition("\n")[2])
        p.add_run("\n")
        document.add_page_break()
        i += 1
    # FIXME: save in a request-specific location such as memory
    document.save('function-security-report.docx')
    headers = {"Content-Disposition": "attachment; filename=%s" % "function-security-report.docx"}
    with open("function-security-report.docx", 'rb') as f:
        body = f.read()
    return make_response((body, headers))


if __name__ == "__main__":
    #Command line options to enable debug and/or saas (bind to 0.0.0.0)
    # print >> sys.stderr, sys.argv
    rand.cleanup()
    csrf_token_raw = rand.bytes(128)
    csrf_token = base64.b64encode(csrf_token_raw)
    port = 5443
    it = iter(sys.argv)
    for arg in it:
        if arg == "--debug":
            # Load default config and override config from an environment variable
            app.config.update(DEBUG=True)
        elif arg == "--saas":
            bindaddr = "0.0.0.0"
        elif arg == "--listen":
            listen = it.next().split(":")
            if len(listen) > 1:
                b, port = listen[:2]
            else:
                b = listen[0]
            if len(b) != 0:
                bindaddr = b
            port = int(port)
    if not os.path.isfile('server.crt'):
       app.run(host=bindaddr, port=port, ssl_context='adhoc')
    else:
       context = SSL.Context(SSL.TLSv1_METHOD)
       context = ('server.crt', 'server.key')
       app.run(host=bindaddr, port=port, ssl_context=context)

