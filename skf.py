# -*- coding: utf-8 -*-
"""
Different Apache license headers for different types of files
# Copyright 2015 Glenn ten Cate, Riccardo ten Cate
Licensed under the Apache License, Version 2.0 (the "License");
you may not use this file except in compliance with the License.
You may obtain a copy of the License at
http://www.apache.org/licenses/LICENSE-2.0
Unless required by applicable law or agreed to in writing, software
distributed under the License is distributed on an "AS IS" BASIS,
WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
See the License for the specific language governing permissions and
limitations under the License.
"""

import os, markdown, datetime
from docx import Document
from BeautifulSoup import BeautifulSoup
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from functools import wraps 
from sqlite3 import dbapi2 as sqlite3
from flask import Flask, request, session, g, redirect, url_for, abort, \
     render_template, flash, Markup, make_response

# create the application
app = Flask(__name__)

def add_response_headers(headers={}):
    """This decorator adds the headers passed in to the response"""
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            resp = make_response(f(*args, **kwargs))
            h = resp.headers
            for header, value in headers.items():
                h[header] = value
            return resp
        return decorated_function
    return decorator

def security(f):
    """This decorator passes multiple security headers"""
    return add_response_headers({'X-Frame-Options': 'deny', 'X-XSS-Protection': '1', 'X-Content-Type-Options': 'nosniff', 'Cache-Control': 'no-store, no-cache', 'Server': 'Security Knowledge Framework'})(f)

# Load default config and override config from an environment variable
app.config.update(dict(
    DATABASE=os.path.join(app.root_path, 'skf.db'),
    DEBUG=True,
    SECRET_KEY='development key',
    USERNAME='admin',
    PASSWORD='default'
))
app.config.from_envvar('skf_SETTINGS', silent=True)


def connect_db():
    """Connects to the specific database."""
    rv = sqlite3.connect(app.config['DATABASE'])
    rv.row_factory = sqlite3.Row
    return rv


def init_db():
    """Initializes the database."""
    db = get_db()
    with app.open_resource('schema.sql', mode='r') as f:
        db.cursor().executescript(f.read())
    db.commit()


@app.cli.command('initdb')
def initdb_command():
    """Creates the database tables."""
    init_db()
    print('Initialized the database.')


def get_db():
    """Opens a new database connection if there is none yet for the
    current application context.
    """
    if not hasattr(g, 'sqlite_db'):
        g.sqlite_db = connect_db()
    return g.sqlite_db

def get_filepaths(directory):
    """
    This function will generate the file names in a directory 
    tree by walking the tree either top-down or bottom-up. For each 
    directory in the tree rooted at directory top (including top itself), 
    it yields a 3-tuple (dirpath, dirnames, filenames).
    """
    file_paths = [] 
    for root, directories, files in os.walk(directory):
        for filename in files:
            filepath = os.path.join(root, filename)
            file_paths.append(filepath)

    return file_paths  

def get_num(x):
    """get numbers from a string"""
    return int(''.join(ele for ele in x if ele.isdigit()))

def get_num_check(x):
    """get numbers from a string for checklists"""
    return [int(s) for s in x.split() if s.isdigit()]

@app.teardown_appcontext
def close_db(error):
    """Closes the database again at the end of the request."""
    if hasattr(g, 'sqlite_db'):
        g.sqlite_db.close()

def projects_functions_techlist():
    """get list of technology used for creating project functions"""
    if not session.get('logged_in'):
        abort(401)
    db = get_db()
    cur = db.execute('SELECT techID, techName, vulnID from techhacks ORDER BY techID DESC')
    entries = cur.fetchall()
    return entries 

@app.route('/')
@security
def show_landing():
    """show the loging page and set default code language"""
    session['code_lang'] = "php"
    return render_template('login.html')

@app.route('/dashboard', methods=['GET'])
@security
def dashboard():
    """show the landing page"""
    if not session.get('logged_in'):
        abort(401)
    return render_template('dashboard.html')

@app.route('/login', methods=['GET', 'POST'])
@security
def login():
    """validate the login data for access dashboard page"""
    error = None
    if request.method == 'POST':
        if request.form['username'] != app.config['USERNAME']:
            error = 'Invalid username/password'
        elif request.form['password'] != app.config['PASSWORD']:
            error = 'Invalid username/password'
        else:
            session['logged_in'] = True
            return render_template('dashboard.html')
    return render_template('login.html', error=error)

@app.route('/logout', methods=['GET', 'POST'])
@security
def logout():
    """logout and destroy session"""
    session.pop('logged_in', None)
    return redirect(url_for('login'))

@app.route('/code/<code_lang>', methods=['GET'])
@security
def set_code_lang(code_lang):
    """set a code language: php java python perl"""
    if not session.get('logged_in'):
        abort(401)
    allowed = "php java python perl"
    found = allowed.find(code_lang)
    if found != -1:
        session['code_lang'] = code_lang
    return redirect(url_for('code_examples'))

@app.route('/code-examples', methods=['GET'])
@security
def code_examples():
    """Shows the knowledge base markdown files."""
    if not session.get('logged_in'):
        abort(401)
    items = []
    id_items = []
    full_file_paths = []
    full_file_paths = get_filepaths(os.path.join(app.root_path, "markdown/code_examples/"+session['code_lang']))
    for path in full_file_paths:
        id_item = get_num(path)
        path = path.split("-")
        y = len(path)-3 
        kb_name_uri = path[(y)]
        kb_name = kb_name_uri.replace("_", " ")
        items.append(kb_name)
        id_items.append(id_item)
    return render_template('code-examples.html', items=items, id_items=id_items)

@app.route('/code-search', methods=['POST'])
@security
def show_code_search():
    """show the landing page"""
    if not session.get('logged_in'):
        abort(401)
    search = request.form['search']
    full_file_paths = []
    full_file_paths = get_filepaths(os.path.join(app.root_path, "markdown/code_examples/"+session['code_lang']))
    for path in full_file_paths:
        found = path.find(search)
        if found != -1:
            filemd = open(path, 'r').read()
            content = Markup(markdown.markdown(filemd))
            path = path.split("-")
            y = len(path)-3
            kb_name_uri = path[(y)]
            kb_name = kb_name_uri.replace("_", " ")
    return render_template('code-examples-search.html', **locals())

@app.route('/code-item', methods=['POST'])
@security
def show_code_item():
    """show the coding examples page"""
    if not session.get('logged_in'):
        abort(401)
    id = int(request.form['id'])
    items = []
    full_file_paths = []
    full_file_paths = get_filepaths(os.path.join(app.root_path, "markdown/code_examples/"+session['code_lang']))
    for path in full_file_paths:
        if id == get_num(path):
            filemd = open(path, 'r').read()
            content = Markup(markdown.markdown(filemd)) 
    return render_template('code-examples-item.html', **locals())

@app.route('/kb-search', methods=['POST'])
@security
def show_kb_search():
    """show the knowledge base search page"""
    if not session.get('logged_in'):
        abort(401)
    search = request.form['search']
    full_file_paths = []
    full_file_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))
    for path in full_file_paths:
        found = path.find(search)
        if found != -1:
            filemd = open(path, 'r').read()
            content = Markup(markdown.markdown(filemd))
            path = path.split("-")
            y = len(path)-3
            kb_name_uri = path[(y)]
            kb_name = kb_name_uri.replace("_", " ")
    return render_template('knowledge-base-search.html', **locals())


@app.route('/kb-item', methods=['POST'])
@security
def show_kb_item():
    """show the knowledge base search result page"""
    if not session.get('logged_in'):
        abort(401)
    id = int(request.form['id'])
    items = []
    full_file_paths = []
    full_file_paths = get_filepaths(os.path.join(app.root_path, "markdown"))
    for path in full_file_paths:
        if id == get_num(path):
            filemd = open(path, 'r').read()
            content = Markup(markdown.markdown(filemd)) 
    return render_template('knowledge-base-item.html', **locals())

@app.route('/knowledge-base', methods=['GET'])
@security
def knowledge_base():
    """Shows the knowledge base markdown files."""
    if not session.get('logged_in'):
        abort(401)
    items = []
    id_items = []
    full_file_paths = []
    full_file_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))
    for path in full_file_paths:
        id_item = get_num(path)
        path = path.split("-")
        y = len(path)-3 
        kb_name_uri = path[(y)]
        kb_name = kb_name_uri.replace("_", " ")
        items.append(kb_name)
        id_items.append(id_item)
    return render_template('knowledge-base.html', items=items, id_items=id_items)

@app.route('/project-new', methods=['GET'])
@security
def projects():
    """show the create new project page"""
    if not session.get('logged_in'):
        abort(401)
    return render_template('project-new.html')

@app.route('/project-add', methods=['POST'])
@security
def add_entry():
    """add a new project to database"""
    if not session.get('logged_in'):
        abort(401)
    db = get_db()
    date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    db.execute('INSERT INTO projects (timestamp, projectName, projectVersion, projectDesc) VALUES (?, ?, ?, ?)',
               [date, request.form['inputName'], request.form['inputVersion'], request.form['inputDesc']])
    db.commit()
    return redirect(url_for('project_list'))

@app.route('/project-del', methods=['POST'])
@security
def project_del():
    """delete project from database"""
    if not session.get('logged_in'):
        abort(401)
    db = get_db()
    db.execute("DELETE from projects WHERE projectID=?",
               [request.form['projectID']])
    db.commit()
    return render_template('reload.html')

@app.route('/project-list', methods=['GET'])
@security
def project_list():
    """show the project list page"""
    if not session.get('logged_in'):
        abort(401)
    db = get_db()
    cur = db.execute('SELECT projectName, projectVersion, projectDESC, projectID, timestamp FROM projects ORDER BY projectID DESC')
    entries = cur.fetchall()
    return render_template('project-list.html', entries=entries)

@app.route('/project-options/<project_id>', methods=['GET'])
@security
def projects_options(project_id):
    """show the project options landing page"""
    if not session.get('logged_in'):
        abort(401)
    return render_template('project-options.html', project_id=project_id)

@app.route('/project-functions/<project_id>', methods=['GET'])
@security
def project_functions(project_id):
    """show the lproject functions page"""
    if not session.get('logged_in'):
        abort(401)
    techlist = projects_functions_techlist()
    db = get_db()
    db.commit()
    cur = db.execute('SELECT paramID, functionName, functionDesc, projectID, tech, entryDate FROM parameters WHERE projectID=? ORDER BY projectID DESC',
                      [project_id])
    entries = cur.fetchall()
    return render_template('project-functions.html', project_id=project_id, techlist=projects_functions_techlist(), entries=entries)

@app.route('/project-function-del', methods=['POST'])
@security
def function_del():
    """delete a project function"""
    if not session.get('logged_in'):
        abort(401)
    id = int(request.form['projectID'])
    db = get_db()
    db.execute("DELETE FROM parameters WHERE projectID=? AND paramID=?",
               [request.form['projectID'],request.form['paramID']])
    db.commit()
    redirect_url = "/project-functions/"+str(id)
    return redirect(redirect_url)


@app.route('/project-function-add', methods=['POST'])
@security
def add_function():
    """add a project function"""
    if not session.get('logged_in'):
        abort(401)
    id = int(request.form['project_id'])
    f = request.form
    for key in f.keys():
        for value in f.getlist(key):
                found = key.find("test")
                if found != -1:
                    db = get_db()
                    date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                    db.execute('INSERT INTO parameters (entryDate, functionName, functionDesc, tech, projectID) VALUES (?, ?, ?, ?, ?)',
                           [date, request.form['functionName'], request.form['functionDesc'], value, request.form['project_id']])
                    db.commit()
    redirect_url = '/project-functions/'+str(id)
    return redirect(redirect_url)

@app.route('/project-checklists/<project_id>', methods=['GET'])
@security
def project_checklists(project_id):
    """show the project checklists page"""
    if not session.get('logged_in'):
        abort(401)
    db = get_db()
    cur = db.execute('SELECT projectName FROM projects WHERE projectID=?',
                        [project_id])
    row = cur.fetchall()
    prep = row[0]
    projectName = prep[0]
    owasp_items = []
    owasp_ids = []
    owasp_kb_ids = []
    owasp_content = []
    custom_items = []
    custom_ids = []
    custom_kb_ids = []
    custom_content = []
    basic_items = []
    basic_ids = []
    basic_kb_ids = []
    basic_content = []
    advanced_items = []
    advanced_ids = []
    advanced_kb_ids = []
    advanced_content = []
    full_file_paths = []
    full_file_paths = get_filepaths(os.path.join(app.root_path, "markdown/checklists"))
    for path in full_file_paths:
       found = path.find("owasp")
       if found != -1:
            owasp_org_path = path
            owasp_list = "owasp"
            owasp_path = path.split("-")
            owasp_kb = owasp_path[5]
            owasp_checklist_name = owasp_path[3]
            owasp_id = get_num(owasp_path[1])
            owasp_items.append(owasp_checklist_name)
            owasp_ids.append(owasp_id)
            owasp_kb_ids.append(owasp_kb)
            filemd = open(owasp_org_path, 'r').read()
            owasp_content.append(Markup(markdown.markdown(filemd)))
    for path in full_file_paths:
       found = path.find("CS_basic_audit")
       if found != -1:
            basic_org_path = path
            basic_list = "CS_basic_audit"
            basic_path = path.split("-")
            basic_kb = basic_path[5]
            basic_checklist_name = basic_path[3]
            basic_id = get_num(basic_path[1])
            basic_items.append(basic_checklist_name)
            basic_ids.append(basic_id)
            basic_kb_ids.append(basic_kb)
            filemd = open(basic_org_path, 'r').read()
            basic_content.append(Markup(markdown.markdown(filemd)))
    for path in full_file_paths:
       found = path.find("CS_advanced_audit")
       if found != -1:
            advanced_org_path = path
            advanced_list = "CS_advanced_audit"
            advanced_path = path.split("-")
            advanced_kb = advanced_path[5]
            advanced_name = advanced_path[3]
            advanced_id = get_num(advanced_path[1])
            advanced_items.append(advanced_name)
            advanced_ids.append(advanced_id)
            advanced_kb_ids.append(advanced_kb)
            filemd = open(advanced_org_path, 'r').read()
            advanced_content.append(Markup(markdown.markdown(filemd)))
    for path in full_file_paths:
       found = path.find("custom")
       if found != -1:
            custom_org_path = path
            custom_list = "custom"
            custom_path = path.split("-")
            custom_kb = custom_path[5]
            custom_name = custom_path[3]
            custom_id = get_num(custom_path[1])
            custom_items.append(custom_name)
            custom_ids.append(custom_id)
            custom_kb_ids.append(custom_kb)
            filemd = open(custom_org_path, 'r').read()
            custom_content.append(Markup(markdown.markdown(filemd)))
    return render_template('project-checklists.html', **locals())

@app.route('/project-checklist-add', methods=['POST'])
@security
def add_checklist():
    """add project checklist"""
    if not session.get('logged_in'):
        abort(401)
    f = request.form
    i = 0
    for key in f.keys():
        for value in f.getlist(key):
            found = key.find("vuln")
            if found != -1:
                listID = "listID"+str(i)
                answerID = "answer"+str(i)
                questionID = "questionID"+str(i) 
                vulnID = "vulnID"+str(i)
                date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                db = get_db()
                db.execute('INSERT INTO questionlist (entryDate, answer, projectName, projectID, questionID, vulnID, listName) VALUES (?, ?, ?, ?, ?, ?, ?)',
                           [date, request.form[answerID], request.form['projectName'], request.form['projectID'], request.form[questionID], request.form[vulnID], request.form[listID]])
                db.commit()
                i += 1
    redirect_url = "/results-checklists"
    return redirect(redirect_url)

@app.route('/results-checklists', methods=['GET'])
@security
def results_checklists():
    """show the results checklists page"""
    if not session.get('logged_in'):
        abort(401)
    db = get_db()
    cur = db.execute('SELECT q.answer, q.projectID, q.questionID,  q.vulnID, q.listName, q.entryDate, p.projectName, p.projectVersion, p.projectDesc FROM questionlist AS q JOIN projects AS p ON q.projectID = p.projectID  GROUP BY q.listName, q.entryDate ORDER BY p.projectName ASC')
    entries = cur.fetchall()
    return render_template('results-checklists.html', entries=entries)

@app.route('/results-functions', methods=['GET'])
@security
def results_functions():
    """show the results functions page"""
    if not session.get('logged_in'):
        abort(401)
    db = get_db()
    cur = db.execute('SELECT p.projectName, p.projectID, p.projectDesc, p.projectVersion, par.paramID, par.functionName, par.projectID FROM projects AS p join parameters AS par on p.projectID = par.projectID GROUP BY p.projectVersion ')
    entries = cur.fetchall()
    return render_template('results-functions.html', entries=entries)

@app.route('/results-functions-del/<entryDate>', methods=['GET'])
@security
def functions_del(entryDate):
    """delete functions result items"""
    if not session.get('logged_in'):
        abort(401)
    db = get_db()
    db.execute("DELETE FROM parameters WHERE entryDate=?",
               [entryDate])
    db.commit()
    redirect_url = "/results-functions"
    return redirect(redirect_url)

@app.route('/results-checklists-del/<entryDate>', methods=['GET'])
@security
def checklists_del(entryDate):
    """delete checklist result item"""
    if not session.get('logged_in'):
        abort(401)
    db = get_db()
    db.execute("DELETE FROM questionlist WHERE entryDate=?",
               [entryDate])
    db.commit()
    redirect_url = "/results-checklists"
    return redirect(redirect_url)


@app.route('/results-checklist-report/<entryDate>', methods=['GET'])
@security
def checklist_results(entryDate):
    """show checklist results report"""
    if not session.get('logged_in'):
        abort(401)
    id_items = []
    content = []
    full_file_paths = []
    db = get_db()
    cur = db.execute("SELECT * FROM questionlist WHERE answer='no' AND entryDate=?",
               [entryDate])
    entries = cur.fetchall()
    for entry in entries:
        projectName = entry[3]
        vulnID = entry[5]
        listName = entry[6]
        entryDate = entry[7]
        full_file_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))
        for path in full_file_paths:
            org_path = path
            path = path.split("-")
            path_vuln = path[1]
            found = path_vuln.find(vulnID)
            if found != -1:
                filemd = open(org_path, 'r').read()
                content.append(Markup(markdown.markdown(filemd)))
    return render_template('results-checklist-report.html', **locals())







@app.route('/results-checklist-docx/<entryDate>')
def download_file(entryDate):
    """Download checklist results report in docx"""
    if not session.get('logged_in'):
        abort(401)
    content_raw = []
    db = get_db()
    cur = db.execute("SELECT * FROM questionlist WHERE answer='no' AND entryDate=?",
               [entryDate])
    entries = cur.fetchall()
    document = Document()
    document.add_picture('static/img/logo.png', width=Inches(1.25))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_heading('Security Knowledge Framework', 0)
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph()
    p.add_run('Used Checklist: ')
    p.add_run('\r\n')
    p.add_run('Date: '+datetime.datetime.now().strftime("%Y-%m-%d %H:%M"))
    p.add_run('\r\n')
    p.add_run('Project: ')
    document.add_page_break()

    p = document.add_heading('Table of contents', level=1)
    p.add_run('\r\n')

    document.add_paragraph('Table of contents', style='IntenseQuote')
    document.add_paragraph('Introduction', style='IntenseQuote')
    for entry in entries:
        projectName = entry[3]
        vulnID = entry[5]
        listName = entry[6]
        entryDate = entry[7]
        full_file_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))


        for path in full_file_paths:
            org_path = path
            path = path.split("-")
            name = org_path[3]
            path_vuln = path[1]
            found = path_vuln.find(vulnID)
            if found != -1:
                filemd = open(org_path, 'r').read()
                content = Markup(markdown.markdown(filemd))
                text = ''.join(BeautifulSoup(content).findAll(text=True))
                text_encode = text.encode('utf-8')
                print text_encode.splitlines()[0]
                content_raw.append(text_encode.splitlines()[0])
                p = document.add_paragraph(text_encode.splitlines()[0], style='IntenseQuote')
                p.add_run()

                #list items found

                #document.add_heading('Introduction', level=1)
                #p = document.add_paragraph(
                #    'The security knowledge framework is composed by means of the highest security standards currently available and is designed to maintain the integrety of your application, so you and your costumers sensitive data is protected against hackers. This document is provided with a checklist in which the programmers of your application had to run through in order to provide a secure product.'
                #)
                #p.add_run('\n')
                #p = document.add_paragraph(
                #    'In the post-development stage of the security knowledge framework the developer double-checks his application against a checklist which consists out of several questions asking the developer about different stages of development and the methodology of implementing different types of functionality the application contains. After filling in this checklist the developer gains feedback on the failed checklist items providing him with solutions about how to solve the additional vulnerability\'s found in the application.'
                #)
                #document.add_page_break()
                #document.add_heading(item.splitlines()[0], level=1)
                #document.add_paragraph(item.partition("\n")[2])
                #document.add_page_break()

    document.save('checklist-security-report.docx')
    headers = {"Content-Disposition": "attachment; filename=%s" % "checklist-security-report.docx"}
    with open("checklist-security-report.docx", 'r') as f:
        body = f.read()
    return make_response((body, headers))
    









